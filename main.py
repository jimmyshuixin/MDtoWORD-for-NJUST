import sys
import os
import shutil
import subprocess
import re
import time
import threading
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from PyQt6.QtWidgets import (QApplication, QMainWindow, QLabel, QVBoxLayout, 
                             QWidget, QProgressBar, QMessageBox, QPushButton, QFileDialog)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QDragEnterEvent, QDropEvent, QCursor
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

"""
NJUST Thesis Formatter
Author: [您的名字]
Repo: https://github.com/your-repo
Description: 将 Markdown 转换为符合 NJUST 规范的 Word 文档。
"""

# 尝试导入 watchdog，如果不存在则使用轮询作为备用
try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    HAS_WATCHDOG = True
except ImportError:
    HAS_WATCHDOG = False

# ==========================================
# 配置与常量：严格映射NJUST规范 
# ==========================================
class NJUST_Config:
    # 纸张与页边距
    PAGE_WIDTH = Mm(210)
    PAGE_HEIGHT = Mm(297)
    MARGIN_TOP = Mm(30)
    MARGIN_BOTTOM = Mm(24)
    MARGIN_LEFT = Mm(25)
    MARGIN_RIGHT = Mm(25)
    HEADER_DIST = Mm(20)
    FOOTER_DIST = Mm(20)

    # 字体名称
    FONT_CN = "SimSun"  # 宋体
    FONT_EN = "Times New Roman"
    FONT_CODE = "Consolas" # 代码块专用字体

    # 字号映射 (Points)
    SIZE_TITLE_1 = Pt(15)   # 小三
    SIZE_TITLE_2 = Pt(14)   # 四号
    SIZE_TITLE_3 = Pt(12)   # 小四
    SIZE_TITLE_4 = Pt(12)   # 小四
    SIZE_BODY = Pt(12)      # 小四
    SIZE_CAPTION = Pt(10.5) # 五号
    SIZE_CODE = Pt(10.5)    # 五号 (代码)
    SIZE_HEADER = Pt(9)     # 小五

    # 间距规则
    LINE_SPACING_BODY = Pt(20) # 固定值20磅

# ==========================================
# 核心逻辑：格式化器
# ==========================================
class NJUST_Formatter:
    def __init__(self, input_path):
        self.input_path = input_path
        self.doc = None 
        
    def setup_page_layout(self):
        """配置页面几何参数"""
        if not self.doc: return
        
        if not self.doc.sections:
            self.doc.add_section()
            
        section = self.doc.sections[0]
        section.page_width = NJUST_Config.PAGE_WIDTH
        section.page_height = NJUST_Config.PAGE_HEIGHT
        section.top_margin = NJUST_Config.MARGIN_TOP
        section.bottom_margin = NJUST_Config.MARGIN_BOTTOM
        section.left_margin = NJUST_Config.MARGIN_LEFT
        section.right_margin = NJUST_Config.MARGIN_RIGHT
        section.header_distance = NJUST_Config.HEADER_DIST
        section.footer_distance = NJUST_Config.FOOTER_DIST
        
        sectPr = section._sectPr
        titlePg = sectPr.get_or_add_titlePg()
        titlePg.val = False 

    def _apply_composite_font(self, run_or_element, size_pt, bold=False, italic=False, force_black=True, is_code=False):
        """应用中西文复合字体"""
        if run_or_element is None: return
        
        is_run_obj = hasattr(run_or_element, 'font')
        rPr = run_or_element._element.get_or_add_rPr() if is_run_obj else run_or_element.get_or_add_rPr()

        if hasattr(size_pt, 'pt'):
            font_size_val = int(size_pt.pt * 2)
        else:
            font_size_val = int(size_pt * 2)

        # 确定使用哪种西文字体
        ascii_font = NJUST_Config.FONT_CODE if is_code else NJUST_Config.FONT_EN

        if is_run_obj:
            run_or_element.font.name = ascii_font
            run_or_element.font.size = Pt(font_size_val / 2)
            run_or_element.font.bold = bold
            run_or_element.font.italic = italic
            if force_black:
                run_or_element.font.color.rgb = RGBColor(0, 0, 0)
                run_or_element.font.underline = False 

        color = rPr.get_or_add_color()
        if force_black:
            color.set(qn('w:val'), '000000')
        
        if force_black:
            u = rPr.get_or_add_u()
            u.set(qn('w:val'), 'none')

        fonts = rPr.get_or_add_rFonts()
        for attr in ['asciiTheme', 'eastAsiaTheme', 'hAnsiTheme', 'cstheme']:
            attr_name = qn('w:' + attr)
            if attr_name in fonts.attrib:
                del fonts.attrib[attr_name]
        
        fonts.set(qn('w:ascii'), ascii_font)
        fonts.set(qn('w:hAnsi'), ascii_font)
        fonts.set(qn('w:eastAsia'), NJUST_Config.FONT_CN)
        fonts.set(qn('w:cs'), ascii_font)
        fonts.set(qn('w:hint'), 'eastAsia') 
        
        sz = rPr.get_or_add_sz()
        sz.set(qn('w:val'), str(font_size_val))
        
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(font_size_val))
        
        if bold:
            bCs = rPr.find(qn('w:bCs'))
            if bCs is None:
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)
            bCs.set(qn('w:val'), '1')

    def _format_paragraph(self, p, level=0):
        """对普通段落应用格式"""
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = NJUST_Config.LINE_SPACING_BODY
        
        if level == 0: 
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.left_indent = Pt(0)
            
            for run in p.runs:
                if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0)
                else:
                    self._apply_composite_font(run, NJUST_Config.SIZE_BODY, bold=False)

        elif level == 1: 
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(18)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE 
            for run in p.runs:
                self._apply_composite_font(run, NJUST_Config.SIZE_TITLE_1, bold=True)

        elif level == 2: 
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in p.runs:
                self._apply_composite_font(run, NJUST_Config.SIZE_TITLE_2, bold=True)
                
        elif level >= 3: 
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in p.runs:
                self._apply_composite_font(run, NJUST_Config.SIZE_TITLE_3, bold=True)

    def _format_reference_paragraph(self, p):
        """参考文献专用格式"""
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = NJUST_Config.LINE_SPACING_BODY
        
        p.paragraph_format.left_indent = Pt(21) 
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        for child in p._element:
            if child.tag.endswith('r'):
                self._apply_composite_font(child, NJUST_Config.SIZE_BODY, bold=False, force_black=True)
            elif child.tag.endswith('hyperlink'):
                for sub_child in child:
                    if sub_child.tag.endswith('r'):
                        self._apply_composite_font(sub_child, NJUST_Config.SIZE_BODY, bold=False, force_black=True)

    def _format_code_block(self, p):
        """[新增] 代码块专用格式"""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE # 代码通常单倍行距
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        # 添加灰色背景 (Shading)
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5') # 浅灰色背景
        pPr.append(shd)
        
        for run in p.runs:
            self._apply_composite_font(run, NJUST_Config.SIZE_CODE, bold=False, is_code=True)

    def _apply_table_style(self, table):
        """应用三线表格式 & 内容居中"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        for child in list(tblBorders):
            tblBorders.remove(child)
            
        def mk_border(name, val, sz, space, color):
            b = OxmlElement(f'w:{name}')
            b.set(qn('w:val'), val)
            b.set(qn('w:sz'), sz) 
            b.set(qn('w:space'), space)
            b.set(qn('w:color'), color)
            return b

        tblBorders.append(mk_border('top', 'single', '12', '0', '000000'))
        tblBorders.append(mk_border('bottom', 'single', '12', '0', '000000'))
        tblBorders.append(mk_border('left', 'nil', '0', '0', 'auto'))
        tblBorders.append(mk_border('right', 'nil', '0', '0', 'auto'))
        tblBorders.append(mk_border('insideV', 'nil', '0', '0', 'auto'))
        tblBorders.append(mk_border('insideH', 'nil', '0', '0', 'auto'))

        table.alignment = WD_TABLE_ALIGNMENT.CENTER 
        
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0) 
                    for run in p.runs:
                        self._apply_composite_font(run, NJUST_Config.SIZE_CAPTION)
                
                if i == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    
                    bottom = tcBorders.find(qn('w:bottom'))
                    if bottom is None:
                        bottom = OxmlElement('w:bottom')
                        tcBorders.append(bottom)
                    
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6') 
                    bottom.set(qn('w:color'), '000000')
                    bottom.set(qn('w:space'), '0')

    def _update_style_font(self, style_name):
        """更新样式定义的默认字体"""
        if style_name in self.doc.styles:
            style = self.doc.styles[style_name]
            if hasattr(style, '_element') and style._element is not None:
                rPr = style._element.get_or_add_rPr()
                fonts = rPr.get_or_add_rFonts()
                fonts.set(qn('w:ascii'), NJUST_Config.FONT_EN)
                fonts.set(qn('w:hAnsi'), NJUST_Config.FONT_EN)
                fonts.set(qn('w:eastAsia'), NJUST_Config.FONT_CN)
                fonts.set(qn('w:cs'), NJUST_Config.FONT_EN)
                    
    def post_process_doc(self, doc):
        """对已有的 Docx 对象进行全量格式清洗"""
        self.doc = doc
        self.setup_page_layout()
        
        for style_id in ['Normal', 'Body Text', 'List Paragraph', 'Heading 1', 'Heading 2', 'Heading 3']:
            self._update_style_font(style_id)
        
        is_reference_section = False
        
        for p in self.doc.paragraphs:
            style_name = p.style.name
            clean_text = p.text.strip().replace(' ', '')
            
            # [新增] 识别 Pandoc 生成的代码块
            if 'Source Code' in style_name or 'Code' in style_name:
                self._format_code_block(p)
                continue

            # 参考文献识别
            if clean_text in ['参考文献', 'References', '参考资料', '主要参考文献']:
                is_reference_section = True
                self._format_paragraph(p, level=1)
                continue
            
            if is_reference_section and clean_text:
                if re.match(r'^\[\d+\]', p.text.strip()) or re.match(r'^\d+\.', p.text.strip()):
                    self._format_reference_paragraph(p)
                else:
                    self._format_paragraph(p, level=0)
                continue

            if style_name.startswith('Heading 1') or style_name.startswith('标题 1'):
                self._format_paragraph(p, level=1)
            elif style_name.startswith('Heading 2') or style_name.startswith('标题 2'):
                self._format_paragraph(p, level=2)
            elif style_name.startswith('Heading 3') or style_name.startswith('标题 3'):
                self._format_paragraph(p, level=3)
            elif style_name.startswith('Caption') or style_name.startswith('Image Caption') or '题注' in style_name:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                for run in p.runs:
                    self._apply_composite_font(run, NJUST_Config.SIZE_CAPTION)
            else:
                self._format_paragraph(p, level=0)
                
        for table in self.doc.tables:
            self._apply_table_style(table)

    def get_safe_output_path(self, base_path):
        """如果文件被占用，自动生成 v1, v2, v3... 后缀"""
        if not os.path.exists(base_path):
            return base_path
            
        try:
            with open(base_path, 'a+'): pass
            return base_path 
        except IOError:
            pass 
            
        folder = os.path.dirname(base_path)
        filename = os.path.basename(base_path)
        name, ext = os.path.splitext(filename)
        
        counter = 1
        while True:
            new_name = f"{name}_v{counter}{ext}"
            new_path = os.path.join(folder, new_name)
            if not os.path.exists(new_path):
                return new_path
            try:
                with open(new_path, 'a+'): pass
                return new_path 
            except IOError:
                counter += 1 

    def convert_with_pandoc(self):
        output_dir = os.path.dirname(self.input_path)
        filename = os.path.basename(self.input_path).rsplit('.', 1)[0]
        temp_docx = os.path.join(output_dir, f"{filename}_temp.docx")
        final_docx = os.path.join(output_dir, f"{filename}_NJUST.docx")
        
        final_docx = self.get_safe_output_path(final_docx)
        
        pandoc_cmd = shutil.which("pandoc")
        if not pandoc_cmd:
            possible_paths = [
                r"C:\Program Files\Pandoc\pandoc.exe",
                r"C:\Program Files (x86)\Pandoc\pandoc.exe",
                os.path.join(os.getenv('LOCALAPPDATA', ''), 'Pandoc', 'pandoc.exe')
            ]
            for p in possible_paths:
                if os.path.exists(p):
                    pandoc_cmd = p
                    break
        
        if not pandoc_cmd:
            raise FileNotFoundError("未找到 Pandoc")

        cmd = [
            pandoc_cmd, 
            self.input_path, 
            '-f', 'markdown+tex_math_dollars+tex_math_single_backslash', 
            '-o', temp_docx,
            '--standalone'
        ]
        
        startupinfo = None
        if os.name == 'nt':
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

        subprocess.run(cmd, check=True, startupinfo=startupinfo)
        
        try:
            doc = Document(temp_docx)
            self.post_process_doc(doc)
            doc.save(final_docx)
        except Exception as e:
            print(f"Post-processing failed: {e}")
            if os.path.exists(temp_docx):
                shutil.copy(temp_docx, final_docx)
            raise e
        finally:
            if os.path.exists(temp_docx):
                try: os.remove(temp_docx)
                except: pass
            
        return final_docx

    def convert_internal(self):
        output_dir = os.path.dirname(self.input_path)
        filename = os.path.basename(self.input_path).rsplit('.', 1)[0]
        output_path = os.path.join(output_dir, f"{filename}_NJUST_Internal.docx")

        output_path = self.get_safe_output_path(output_path)

        self.doc = Document()
        self.setup_page_layout()
        
        with open(self.input_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
            
        html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup:
            if isinstance(element, NavigableString):
                if element.strip(): self.add_paragraph_internal(element.strip())
                continue
                
            if element.name == 'h1': self.add_heading_internal(element.text, 1)
            elif element.name == 'h2': self.add_heading_internal(element.text, 2)
            elif element.name == 'h3': self.add_heading_internal(element.text, 3)
            # [新增] 识别 pre 代码块
            elif element.name == 'pre': 
                self.add_code_block_internal(element)
            elif element.name == 'p':
                img = element.find('img')
                if img and len(element.get_text(strip=True)) == 0:
                    self.add_image_internal(img['src'], img.get('alt', ''))
                else:
                    self.add_rich_paragraph_internal(element)
            elif element.name == 'table': self.add_table_internal(element)
            elif element.name in ['ul', 'ol']: self.add_list_internal(element, element.name=='ol')

        self.doc.save(output_path)
        return output_path

    # ... (Add methods) ...
    def add_heading_internal(self, text, level):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._format_paragraph(p, level=level)

    def add_paragraph_internal(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._format_paragraph(p, level=0)

    def add_code_block_internal(self, element):
        """[新增] 内置引擎处理代码块"""
        text = element.get_text()
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._format_code_block(p) # 应用代码块样式

    def add_rich_paragraph_internal(self, soup_element):
        p = self.doc.add_paragraph()
        self._format_paragraph(p, level=0)
        for child in soup_element.contents:
            if isinstance(child, NavigableString):
                text = str(child)
                if text: 
                    run = p.add_run(text)
                    self._apply_composite_font(run, NJUST_Config.SIZE_BODY)
            elif isinstance(child, Tag):
                text = child.get_text()
                is_bold = child.name in ['strong', 'b']
                is_italic = child.name in ['em', 'i']
                # 处理行内代码 `code`
                is_code = child.name == 'code'
                run = p.add_run(text)
                self._apply_composite_font(run, NJUST_Config.SIZE_BODY, bold=is_bold, italic=is_italic, is_code=is_code)

    def add_image_internal(self, src, caption):
        if not os.path.isabs(src):
            src = os.path.join(os.path.dirname(self.input_path), src)
        if os.path.exists(src):
            try:
                self.doc.add_picture(src, width=Mm(160))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"图 {caption}")
                    self._apply_composite_font(run, NJUST_Config.SIZE_CAPTION)
            except: pass

    def add_table_internal(self, table_element):
        rows = table_element.find_all('tr')
        if not rows: return
        max_cols = max([len(r.find_all(['td', 'th'])) for r in rows])
        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        for i, row in enumerate(rows):
            cols = row.find_all(['td', 'th'])
            for j, col in enumerate(cols):
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(col.get_text(strip=True))
        self._apply_table_style(table)

    def add_list_internal(self, element, ordered=False):
        for i, li in enumerate(element.find_all('li', recursive=False)):
            text = li.get_text(strip=True)
            prefix = f"{i+1}. " if ordered else "● "
            p = self.doc.add_paragraph()
            self._format_paragraph(p, level=0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.first_line_indent = Pt(-21)
            run = p.add_run(prefix + text)
            self._apply_composite_font(run, NJUST_Config.SIZE_BODY)

# ==========================================
# 文件夹监控线程 (使用 watchdog)
# ==========================================
class WatchdogWorker(QThread):
    """
    [改进] 使用 Watchdog 实现的高效文件监控器
    """
    file_detected_signal = pyqtSignal(str)

    def __init__(self, folder_path):
        super().__init__()
        self.folder_path = folder_path
        self.observer = None
        self.handler = None

    def run(self):
        if not HAS_WATCHDOG:
            print("Watchdog not found, falling back to polling.")
            self._run_polling()
            return

        # 创建 Watchdog 事件处理器
        class NewFileHandler(FileSystemEventHandler):
            def __init__(self, signal):
                self.signal = signal
                self.last_processed = {} # 防止重复处理 (文件创建+修改可能触发多次)

            def on_created(self, event):
                if not event.is_directory and event.src_path.lower().endswith('.md'):
                    self._trigger(event.src_path)

            def on_moved(self, event):
                if not event.is_directory and event.dest_path.lower().endswith('.md'):
                    self._trigger(event.dest_path)
            
            def _trigger(self, path):
                now = time.time()
                # 简单防抖：3秒内不重复处理同一文件
                if path in self.last_processed and (now - self.last_processed[path] < 3):
                    return
                self.last_processed[path] = now
                
                # 等待文件写入释放 (防止文件刚创建还没内容)
                time.sleep(0.5) 
                self.signal.emit(path)

        self.handler = NewFileHandler(self.file_detected_signal)
        self.observer = Observer()
        self.observer.schedule(self.handler, self.folder_path, recursive=False)
        self.observer.start()
        
        # 保持线程运行
        try:
            while self.observer.is_alive():
                self.observer.join(1)
        except:
            self.observer.stop()
        self.observer.join()

    def _run_polling(self):
        """Watchdog 缺失时的备用方案"""
        known_files = set()
        if os.path.exists(self.folder_path):
            for f in os.listdir(self.folder_path):
                if f.lower().endswith('.md'):
                    known_files.add(f)
        
        while True:
            if self.isInterruptionRequested(): break
            if not os.path.exists(self.folder_path):
                time.sleep(2)
                continue
                
            current_files = set()
            for f in os.listdir(self.folder_path):
                if f.lower().endswith('.md'):
                    current_files.add(f)
                    if f not in known_files:
                        time.sleep(1)
                        known_files.add(f)
                        self.file_detected_signal.emit(os.path.join(self.folder_path, f))
            known_files = current_files
            time.sleep(2)

    def stop(self):
        if self.observer:
            self.observer.stop()
        self.requestInterruption()

# ==========================================
# 转换工作线程
# ==========================================
class WorkerThread(QThread):
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    info_signal = pyqtSignal(str) 

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:
            formatter = NJUST_Formatter(self.file_path)
            
            try:
                self.info_signal.emit("正在尝试使用 Pandoc 引擎...")
                output_path = formatter.convert_with_pandoc()
                self.finished_signal.emit(output_path)
                return
            except PermissionError as e:
                self.error_signal.emit(str(e))
                return
            except FileNotFoundError:
                self.info_signal.emit("未检测到 Pandoc，切换至内置引擎...")
            except Exception as e:
                print(f"Pandoc error: {e}") 
                self.info_signal.emit(f"Pandoc 转换出错，切换至内置引擎...")

            self.info_signal.emit("正在使用内置引擎解析...")
            output_path = formatter.convert_internal()
            self.finished_signal.emit(output_path)
            
        except Exception as e:
            import traceback
            self.error_signal.emit(str(e) + "\n" + traceback.format_exc())

# ==========================================
# 主窗口
# ==========================================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("NJUST 论文格式转换工具 (自动监控版)")
        self.resize(600, 600)
        self.setAcceptDrops(True)
        self.watcher_thread = None
        self.init_ui()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 拖拽区域
        self.label = QLabel("模式一：将 Markdown (.md) 文件拖入此处")
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label.setStyleSheet("""
            QLabel {
                border: 3px dashed #aaa;
                border-radius: 10px;
                font-size: 16px;
                color: #555;
                background-color: #f9f9f9;
                padding: 30px;
            }
        """)
        layout.addWidget(self.label)
        
        # 监控按钮
        self.monitor_btn = QPushButton("模式二：选择并监控文件夹 (自动转换)")
        self.monitor_btn.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.monitor_btn.setMinimumHeight(50)
        self.monitor_btn.setStyleSheet("""
            QPushButton {
                background-color: #0078D7;
                color: white;
                font-size: 14px;
                font-weight: bold;
                border-radius: 5px;
                border: none;
            }
            QPushButton:hover { background-color: #005A9E; }
            QPushButton:pressed { background-color: #004578; }
        """)
        self.monitor_btn.clicked.connect(self.select_folder)
        layout.addWidget(self.monitor_btn)
        
        self.monitor_label = QLabel("当前未监控任何文件夹")
        self.monitor_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.monitor_label.setStyleSheet("color: #666; font-size: 12px;")
        layout.addWidget(self.monitor_label)

        self.status_label = QLabel("就绪")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setStyleSheet("font-size: 13px; color: #333;")
        layout.addWidget(self.status_label)

        self.progress = QProgressBar()
        self.progress.setVisible(False)
        self.progress.setStyleSheet("""
            QProgressBar { height: 6px; border-radius: 3px; background: #eee; }
            QProgressBar::chunk { background-color: #0078D7; border-radius: 3px; }
        """)
        layout.addWidget(self.progress)
        
        # 底部信息栏布局
        bottom_layout = QVBoxLayout()
        bottom_layout.setSpacing(5)

        # 依赖检测提示
        ver_info = "V3.2 | 代码块识别 | "
        ver_info += "Watchdog 监控中" if HAS_WATCHDOG else "Watchdog 未安装 (轮询模式)"
        version_label = QLabel(ver_info)
        version_label.setAlignment(Qt.AlignmentFlag.AlignRight)
        version_label.setStyleSheet("color: #999; font-size: 10px;")
        bottom_layout.addWidget(version_label)

        # =========================================
        # [新增] 作者信息与引导链接
        # =========================================
        # 请修改下方的 href 和 文本内容
        author_text = (
            'Created by <a href="https://github.com/jimmyshuixin/MDtoWORD-for-NJUST" style="color:#0078D7; text-decoration:none;">'
            '[JimmyShuixin]</a> | '
            '<a href="https://github.com/jimmyshuixin/MDtoWORD-for-NJUST/blob/main/README.md" style="color:#0078D7; text-decoration:none;">'
            '查看使用教程 & 帮助</a>'
        )
        self.author_label = QLabel(author_text)
        self.author_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.author_label.setOpenExternalLinks(True) # 允许点击跳转浏览器
        self.author_label.setStyleSheet("QLabel { color: #666; font-size: 11px; margin-top: 5px; }")
        
        bottom_layout.addWidget(self.author_label)
        layout.addLayout(bottom_layout)

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择要监控的 Markdown 文件夹")
        if folder:
            if self.watcher_thread and self.watcher_thread.isRunning():
                self.watcher_thread.stop()
                self.watcher_thread.wait()
            
            self.monitor_label.setText(f"正在监控: {folder}\n(将自动转换新增的 .md 文件)")
            self.monitor_label.setStyleSheet("color: #2E7D32; font-weight: bold;")
            
            # 启动新监控线程 (WatchdogWorker)
            self.watcher_thread = WatchdogWorker(folder)
            self.watcher_thread.file_detected_signal.connect(self.start_conversion_silent)
            self.watcher_thread.start()
            
            msg = f"已开始监控文件夹：\n{folder}\n\n"
            if not HAS_WATCHDOG:
                msg += "⚠️ 提示：未检测到 watchdog 库，当前使用轮询模式 (Polling)。\n建议安装: pip install watchdog 以获得更好体验。"
            
            QMessageBox.information(self, "监控已启动", msg)

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if urls and urls[0].toLocalFile().lower().endswith('.md'):
                event.accept()
                self.label.setStyleSheet("QLabel { border: 3px dashed #4CAF50; background-color: #E8F5E9; color: #2E7D32; font-size: 16px; padding: 30px; }")
                self.label.setText("释放以开始转换")
            else:
                event.ignore()
        else:
            event.ignore()

    def dragLeaveEvent(self, event):
        self.label.setText("模式一：将 Markdown (.md) 文件拖入此处")
        self.label.setStyleSheet("QLabel { border: 3px dashed #aaa; background-color: #f9f9f9; font-size: 16px; color: #555; padding: 30px; }")

    def dropEvent(self, event: QDropEvent):
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if file_path.lower().endswith('.md'):
                self.start_conversion(file_path)

    def start_conversion(self, file_path):
        self.label.setText(f"处理中：{os.path.basename(file_path)}")
        self._run_conversion_worker(file_path)

    def start_conversion_silent(self, file_path):
        self.status_label.setText(f"检测到新文件：{os.path.basename(file_path)}")
        self._run_conversion_worker(file_path)

    def _run_conversion_worker(self, file_path):
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        
        self.worker = WorkerThread(file_path)
        self.worker.finished_signal.connect(self.on_success)
        self.worker.error_signal.connect(self.on_error)
        self.worker.info_signal.connect(self.update_status)
        self.worker.start()

    def update_status(self, msg):
        self.status_label.setText(msg)

    def on_success(self, output_path):
        self.progress.setVisible(False)
        self.label.setText("转换成功！")
        self.status_label.setText(f"已生成: {os.path.basename(output_path)}")
        self.label.setStyleSheet("QLabel { border: 3px solid #4CAF50; color: #4CAF50; font-size: 16px; padding: 30px; }")
        
        try:
            if sys.platform == 'win32':
                os.startfile(output_path)
            elif sys.platform == 'darwin':
                subprocess.call(('open', output_path))
            else:
                subprocess.call(('xdg-open', output_path))
        except:
            pass

    def on_error(self, err_msg):
        self.progress.setVisible(False)
        self.label.setText("转换出错")
        self.label.setStyleSheet("QLabel { border: 3px solid #F44336; color: #F44336; font-size: 16px; padding: 30px; }")
        self.status_label.setText(f"错误: {err_msg[:50]}...")

    def closeEvent(self, event):
        if self.watcher_thread:
            self.watcher_thread.stop()
            self.watcher_thread.wait()
        event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
